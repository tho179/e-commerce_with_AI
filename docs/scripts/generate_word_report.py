from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[2]
DOCS_DIR = PROJECT_ROOT / "docs"
OUTPUT_PATH = DOCS_DIR / "bao-cao-AI-services-15-20-trang.docx"

DATA_CSV = PROJECT_ROOT / "ai-service" / "data" / "data_user500.csv"
METRICS_CSV = PROJECT_ROOT / "ai-service" / "data" / "ml_outputs" / "metrics_summary.csv"
TRAIN_CODE = PROJECT_ROOT / "ai-service" / "models" / "train_sequence_models.py"
GRAPH_CODE = PROJECT_ROOT / "ai-service" / "graph_logic" / "build_kb_graph.py"

IMAGES = {
    "data_top20": DOCS_DIR / "images" / "data_user500_top20.png",
    "metric_cmp": PROJECT_ROOT / "ai-service" / "data" / "ml_outputs" / "metric_comparison.png",
    "curve_rnn": PROJECT_ROOT / "ai-service" / "data" / "ml_outputs" / "curve_RNN.png",
    "curve_lstm": PROJECT_ROOT / "ai-service" / "data" / "ml_outputs" / "curve_LSTM.png",
    "curve_bilstm": PROJECT_ROOT / "ai-service" / "data" / "ml_outputs" / "curve_biLSTM.png",
    "cm_best": PROJECT_ROOT / "ai-service" / "data" / "ml_outputs" / "confusion_matrix_best.png",
    "kb_complex": DOCS_DIR / "images" / "kb_graph_complex.png",
    "kb_rows": DOCS_DIR / "images" / "kb_graph_top20_rows.png",
    "rag_pipeline": DOCS_DIR / "images" / "rag_pipeline_2d.png",
    "scene_full": DOCS_DIR / "images" / "scene_full_structure.png",
    "scene_user": DOCS_DIR / "images" / "scene_user_centric.png",
    "scene_20": DOCS_DIR / "images" / "scene_graph_20_rows.png",
    "scene_query": DOCS_DIR / "images" / "scene_neo4j_query.png",
}


def set_default_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(13)

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    if "CodeBlock" not in [s.name for s in doc.styles]:
        style = doc.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Consolas"
        style.font.size = Pt(8.5)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TRƯỜNG / KHOA: ......................................................")
    run.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("BÁO CÁO ĐỒ ÁN AI SERVICE").bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Tích hợp AISERVICE, KB_Graph và Graph-RAG Chat vào hệ e-commerce").bold = True

    doc.add_paragraph("")
    doc.add_paragraph("Môn học: Hệ thống thông tin / Thương mại điện tử")
    doc.add_paragraph("Phạm vi: Câu 2a, 2c, 2d và triển khai thực nghiệm")
    doc.add_paragraph("Sinh viên: ..........................................................")
    doc.add_paragraph("MSSV: ..........................................................")
    doc.add_paragraph("Lớp: ..........................................................")
    doc.add_paragraph("Giảng viên hướng dẫn: ..........................................................")

    p = doc.add_paragraph("TP.HCM, ngày 21 tháng 04 năm 2026")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraphs(doc: Document, paragraphs: list[str]) -> None:
    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.35


def add_image(doc: Document, path: Path, caption: str, width_cm: float = 15.5) -> None:
    if not path.exists():
        doc.add_paragraph(f"[Thiếu ảnh: {path.name}]")
        return

    doc.add_picture(str(path), width=Cm(width_cm))
    pic = doc.paragraphs[-1]
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True


def add_dataframe_table(doc: Document, title: str, df: pd.DataFrame) -> None:
    doc.add_paragraph(title)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    for idx, col in enumerate(df.columns):
        cell = table.rows[0].cells[idx]
        cell.text = str(col)
        for run in cell.paragraphs[0].runs:
            run.bold = True

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, val in enumerate(row.tolist()):
            cells[idx].text = str(val)


def add_code_block(doc: Document, title: str, code_text: str) -> None:
    doc.add_paragraph(title)
    lines = code_text.splitlines()

    chunk_size = 90
    for i in range(0, len(lines), chunk_size):
        chunk = "\n".join(lines[i : i + chunk_size])
        p = doc.add_paragraph(style="CodeBlock")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.0
        p.add_run(chunk)


def build_document() -> None:
    doc = Document()
    set_default_style(doc)
    add_title_page(doc)

    add_heading(doc, "TÓM TẮT BÁO CÁO", level=1)
    add_paragraphs(
        doc,
        [
            "Báo cáo này trình bày lộ trình triển khai đầy đủ từ mô hình dự đoán hành vi người dùng (Câu 2a), xây dựng tri thức đồ thị KB_Graph (Câu 2c), đến tích hợp Graph-RAG cho chatbot tư vấn trên hệ e-commerce (Câu 2d).",
            "Mục tiêu của nhóm giải pháp là chuyển AI từ mức phân tích đơn lẻ sang mức hỗ trợ quyết định trong thời gian thực: vừa tư vấn sản phẩm, vừa cung cấp bằng chứng truy xuất, đồng thời giữ khả năng fallback để đảm bảo độ ổn định vận hành.",
            "Về mặt kiến trúc, AISERVICE đóng vai trò trung tâm, gom semantic search, recommendation, review intelligence, drift monitor và Graph-RAG vào cùng một service. Điều này giảm mạnh độ phân mảnh và chi phí vận hành so với kiến trúc nhiều AI service rời rạc.",
            "Bản Word này được mở rộng theo chuẩn nộp học thuật (15-20 trang), gồm nội dung phân tích, bảng dữ liệu, hình minh họa scene, kết quả định lượng, và phụ lục mã nguồn quan trọng để hội đồng có thể kiểm tra đầy đủ mức độ hoàn thiện kỹ thuật.",
        ],
    )

    add_heading(doc, "1. MÔ TẢ AISERVICE", level=1)
    add_paragraphs(
        doc,
        [
            "AISERVICE là service AI trung tâm trong hệ thống, chịu trách nhiệm xử lý các năng lực cốt lõi: semantic search, recommendation, review intelligence, drift monitor, retrain trigger và chatbot. Việc hợp nhất này giúp API Gateway chỉ cần gọi một điểm tích hợp duy nhất.",
            "Với mô hình hợp nhất, các endpoint có thể chia sẻ dữ liệu nội bộ hiệu quả hơn, đặc biệt là khi cần kết hợp nhiều nguồn ngữ cảnh (hành vi user, metadata sản phẩm, insight đánh giá) để trả lời câu hỏi nghiệp vụ.",
            "Ở tầng chatbot, hệ thống hỗ trợ hai luồng: ChatAdviceView cho semantic/policy thông thường và GraphRAGChatView cho truy xuất tri thức từ Neo4j. Luồng GraphRAG có cơ chế fallback semantic nhằm giảm nguy cơ lỗi gián đoạn dịch vụ.",
            "Từ góc nhìn vận hành, AISERVICE phù hợp với microservices e-commerce vì dễ quan sát trạng thái, dễ test tích hợp và thuận tiện rollback theo một đơn vị triển khai thống nhất.",
        ],
    )

    add_heading(doc, "2. DỮ LIỆU ĐẦU VÀO VÀ TIỀN XỬ LÝ", level=1)
    data_df = pd.read_csv(DATA_CSV).head(20)
    add_paragraphs(
        doc,
        [
            "Nguồn dữ liệu hành vi gồm các trường user_id, product_id, action, timestamp. Tập dữ liệu phản ánh chuỗi tương tác điển hình của người dùng trong môi trường mua sắm số: xem, click, tìm kiếm, thêm yêu thích, thêm giỏ, checkout, purchase.",
            "Trong bước tiền xử lý cho bài toán sequence learning, timestamp được chuyển thành đặc trưng chu kỳ thời gian (hour_sin, hour_cos, dow_sin, dow_cos), giúp mô hình học được nhịp hành vi theo khung giờ và ngày trong tuần.",
        ],
    )
    add_dataframe_table(doc, "Bảng 1. Hai mươi dòng đầu của data_user500.csv", data_df)
    add_image(doc, IMAGES["data_top20"], "Hình 1. Ảnh minh họa 20 dòng dữ liệu đầu vào", width_cm=16.0)

    add_heading(doc, "3. CÂU 2a - MÔ HÌNH RNN/LSTM/biLSTM", level=1)
    add_paragraphs(
        doc,
        [
            "Câu 2a thiết lập bài toán dự đoán hành động kế tiếp bằng mô hình tuần tự. Input mỗi mẫu gồm chuỗi độ dài cố định (SEQ_LEN=3) của action/product/time, target là action tiếp theo.",
            "Ba kiến trúc được đánh giá đồng thời: RNN cơ bản, LSTM một chiều và biLSTM hai chiều. Các mô hình dùng embedding cho action và product, sau đó nối với vector thời gian để đưa vào khối recurrent.",
            "Chiến lược huấn luyện có early stopping dựa trên macro_f1 của tập validation, nhằm hạn chế overfit và ưu tiên chất lượng phân lớp đa lớp cân bằng hơn so với chỉ tối ưu accuracy.",
            "Kết quả thực nghiệm cho thấy LSTM có macro_f1 cao nhất trong ba mô hình, vì vậy được chọn làm model_best cho các bước tích hợp tiếp theo.",
        ],
    )

    metrics_df = pd.read_csv(METRICS_CSV)
    add_dataframe_table(doc, "Bảng 2. Kết quả so sánh mô hình (test)", metrics_df)
    add_image(doc, IMAGES["curve_rnn"], "Hình 2. Đường cong huấn luyện RNN", width_cm=14.2)
    add_image(doc, IMAGES["curve_lstm"], "Hình 3. Đường cong huấn luyện LSTM", width_cm=14.2)
    add_image(doc, IMAGES["curve_bilstm"], "Hình 4. Đường cong huấn luyện biLSTM", width_cm=14.2)
    add_image(doc, IMAGES["metric_cmp"], "Hình 5. So sánh các chỉ số đánh giá", width_cm=15.0)
    add_image(doc, IMAGES["cm_best"], "Hình 6. Confusion matrix của model_best", width_cm=13.2)

    add_heading(doc, "4. CÂU 2c - XÂY DỰNG KB_GRAPH", level=1)
    add_paragraphs(
        doc,
        [
            "Ở Câu 2c, dữ liệu log được ánh xạ thành đồ thị tri thức gồm các node User, Product, Action, Event, TimeSlot và quan hệ liên kết thể hiện ngữ nghĩa hành vi. Cấu trúc này cho phép truy vấn phong phú hơn bảng quan hệ truyền thống.",
            "Pipeline import thực hiện theo batch: đọc CSV, chuẩn hóa thời gian, sinh event_id, merge node/edge vào Neo4j, sau đó tổng hợp các quan hệ ở mức cao như INTERACTED_WITH và NEXT_ACTION.",
            "Nhờ biểu diễn đồ thị, hệ thống có thể trả lời nhanh các câu hỏi như: sản phẩm nào có funnel tốt, hành vi chuyển tiếp phổ biến, hay khung giờ tương tác cao. Đây là nền tảng trực tiếp cho Graph-RAG ở Câu 2d.",
        ],
    )

    add_image(doc, IMAGES["kb_rows"], "Hình 7. Top 20 dòng tổng hợp cho KB_Graph", width_cm=15.2)
    add_image(doc, IMAGES["kb_complex"], "Hình 8. KB_Graph mật độ cao nhiều lớp", width_cm=16.5)

    add_heading(doc, "5. CÂU 2d - GRAPH-RAG CHAT VÀ TÍCH HỢP E-COMMERCE", level=1)
    add_paragraphs(
        doc,
        [
            "Câu 2d tập trung đưa tri thức đồ thị vào hội thoại tư vấn. Luồng GraphRAGChatView nhận query, phát hiện intent, truy vấn Neo4j qua KBGraphRAG, tổng hợp câu trả lời có citations và graph_context rồi trả về UI.",
            "Khi Neo4j tạm thời unavailable hoặc kết quả graph không đủ thông tin, service fallback về semantic search để đảm bảo trải nghiệm không bị ngắt quãng. Đây là cơ chế quan trọng trong triển khai thực tế.",
            "Dữ liệu đầu ra không chỉ là answer dạng text mà còn bao gồm products, recommendations, confidence và suggested_prompts. Nhờ vậy frontend có thể render giao diện chat concierge giàu ngữ cảnh thay vì chat text đơn thuần.",
            "Điểm khác biệt của hướng Graph-RAG là khả năng truy xuất có căn cứ. Người dùng và đội vận hành đều có thể nhìn thấy nguồn tri thức, qua đó giảm hallucination và tăng tính kiểm chứng của phản hồi.",
        ],
    )
    add_image(doc, IMAGES["rag_pipeline"], "Hình 9. Pipeline Graph-RAG chatbot", width_cm=16.0)

    add_heading(doc, "6. SCENE ĐỒ THỊ BỔ SUNG", level=1)
    add_image(doc, IMAGES["scene_full"], "Hình 10. Scene toàn bộ cấu trúc", width_cm=16.8)
    add_image(doc, IMAGES["scene_user"], "Hình 11. Scene user-centric", width_cm=14.8)
    add_image(doc, IMAGES["scene_20"], "Hình 12. Scene 20 dòng Graph", width_cm=15.8)
    add_image(doc, IMAGES["scene_query"], "Hình 13. Scene truy vấn Neo4j", width_cm=16.0)

    add_heading(doc, "7. ĐÁNH GIÁ, RỦI RO VÀ HƯỚNG PHÁT TRIỂN", level=1)
    add_paragraphs(
        doc,
        [
            "Ưu điểm chính của kiến trúc hiện tại là tính mô-đun nhưng vẫn thống nhất ở tầng AI service. Điều này giúp mở rộng tính năng nhanh và giảm độ phân mảnh trong bảo trì.",
            "Rủi ro kỹ thuật còn tồn tại gồm: chất lượng dữ liệu hành vi chưa đồng đều giữa các ngành hàng, độ bao phủ truy vấn graph phụ thuộc vào mức độ đầy đủ dữ liệu event, và chi phí truy vấn tăng theo mật độ đồ thị.",
            "Trong giai đoạn tiếp theo, nhóm nên bổ sung cơ chế đánh giá online A/B cho chatbot, giám sát mô hình theo thời gian thực, và tái huấn luyện định kỳ dựa trên drift trigger có kiểm soát.",
            "Về nghiệp vụ, cần đo tác động của Graph-RAG đến KPI chuyển đổi (CTR, add-to-cart rate, checkout rate, purchase rate) để chứng minh giá trị kinh doanh một cách định lượng.",
            "Ngoài ra, có thể tích hợp thêm vector retrieval cho mô tả sản phẩm dài, sau đó hợp nhất với graph retrieval thành hybrid-RAG để nâng chất lượng câu trả lời cho truy vấn phức tạp.",
        ],
    )

    doc.add_page_break()

    add_heading(doc, "PHỤ LỤC A - MÃ NGUỒN CÂU 2a (TOÀN VĂN)", level=1)
    train_code_text = TRAIN_CODE.read_text(encoding="utf-8") if TRAIN_CODE.exists() else "Không tìm thấy file mã nguồn."
    add_code_block(doc, "A.1 File ml/train_sequence_models.py", train_code_text)

    doc.add_page_break()

    add_heading(doc, "PHỤ LỤC B - MÃ NGUỒN BUILD KB_GRAPH", level=1)
    graph_code_text = GRAPH_CODE.read_text(encoding="utf-8") if GRAPH_CODE.exists() else "Không tìm thấy file mã nguồn."
    add_code_block(doc, "B.1 File kb_graph/build_kb_graph.py", graph_code_text)

    doc.add_page_break()

    add_heading(doc, "PHỤ LỤC C - MẪU TRUY VẤN NEO4J", level=1)
    cypher_queries = """-- Top sản phẩm theo purchases
MATCH (:User)-[r:INTERACTED_WITH]->(p:Product)
RETURN p.product_id AS product_id, sum(r.purchases) AS purchases
ORDER BY purchases DESC
LIMIT 10;

-- Chuỗi hành vi phổ biến
MATCH (a1:Action)-[r:NEXT_ACTION]->(a2:Action)
RETURN a1.name AS from_action, a2.name AS to_action, r.count AS transition_count
ORDER BY transition_count DESC
LIMIT 20;

-- Khung giờ tương tác cao
MATCH (e:Event)-[:IN_TIMESLOT]->(t:TimeSlot)
RETURN t.date AS date, t.hour AS hour, count(e) AS events
ORDER BY events DESC
LIMIT 20;

-- Funnel của một sản phẩm cụ thể
MATCH (:User)-[r:INTERACTED_WITH]->(p:Product {product_id: 'P0001'})
RETURN p.product_id,
       sum(r.views) AS views,
       sum(r.cart_adds) AS cart_adds,
       sum(r.checkouts) AS checkouts,
       sum(r.purchases) AS purchases;
"""
    add_code_block(doc, "C.1 Bộ truy vấn tham chiếu", cypher_queries)

    add_heading(doc, "KẾT LUẬN", level=1)
    add_paragraphs(
        doc,
        [
            "Tài liệu đã trình bày đầy đủ chuỗi giá trị AI từ dữ liệu hành vi -> mô hình dự đoán -> tri thức đồ thị -> Graph-RAG chat -> tích hợp UI e-commerce. Việc kết hợp này cho thấy khả năng triển khai thực tế, có tính mở rộng và có thể đánh giá định lượng theo KPI vận hành và KPI kinh doanh.",
            "Phiên bản Word được mở rộng với mục tiêu độ dài 15-20 trang để phục vụ nộp báo cáo chính thức, đồng thời giữ tính kiểm chứng nhờ hình ảnh thực nghiệm và phụ lục mã nguồn đi kèm.",
        ],
    )

    doc.save(OUTPUT_PATH)
    print(f"Created Word report: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
